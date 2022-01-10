#!/usr/bin/python3

import argparse
import datetime
import json
import logging
import os
import pandas as pd
import plotly.io as pio
import re
import requests
import subprocess
import sys
import tarfile
from sqlalchemy import create_engine
from docx import Document
from docx.shared import Inches
from elasticsearch import Elasticsearch, helpers


# Configure Logging
logging.basicConfig(stream=sys.stdout, level=logging.INFO)
logger = logging.getLogger(__name__)


def generate_graphs(directory):
    graphs = {}
    # Walking through the folder to find all the result files
    # ignoring dirname since we don't need it in this implementation
    for root, _, files in os.walk(directory):
        for filename in files:
            regex = re.compile(r'[\w-]+_([\w-]+).(\w.+)')
            matches = regex.match(filename)

            if 'summaries' not in root and regex.match(filename) and \
                    matches.group(2) == 'json':
                logger.info(f'Generating graph for: {matches.group(1)}')

                # Initializes database for current file in current directory
                # Read by 20000 chunks
                disk_engine = create_engine(
                    'sqlite:///{}/{}.db'.format(root, matches.group(1)))

                j = 0
                index_start = 1
                chunk = 20000
                for df in pd.read_json(root+'/'+filename,
                                       lines=True,
                                       chunksize=chunk):
                    df.index += index_start

                    columns = ['timestamp', 'latency', 'code']

                    for c in df.columns:
                        if c not in columns:
                            df = df.drop(c, axis=1)

                    j += 1
                    logger.info(f'completed {j*chunk} rows')

                    df.to_sql('data', disk_engine, if_exists='append')
                    index_start = df.index[-1] + 1

                df = pd.read_sql_query('SELECT * FROM data', disk_engine)

                data = [{
                    'type': 'scatter',
                    'x': df['timestamp'],
                    'y': df['latency']/1000000,
                    'mode': 'markers',
                    'transforms': [{
                        'type': 'groupby',
                        'groups': df['code'],
                        'color': 'code',
                        'styles': [
                            {'target': '0',
                             'value': {'marker': {'color': 'coral',
                                                  'symbol': 'triangle-down'}}},
                            {'target': '200',
                             'value': {'marker': {'color': 'LightSkyBlue'}}},
                            {'target': '201',
                             'value': {'marker': {'color': 'LightSkyBlue'}}},
                            {'target': '400',
                             'value': {'marker': {'color': 'crimsom',
                                                  'symbol': 'diamond'}}},
                            {'target': '500',
                             'value': {'marker': {'color': 'darkred',
                                                  'symbol': 'diamond-tall'}}}]
                    }]
                }]

                layout = {
                    'title': '<b>Latency per Request: {}</b>'.format(
                        matches.group(1)),
                    'xaxis': {'title': 'Time',
                              'showgrid': 'true',
                              'ticklabelmode': "period"},
                    'yaxis': {'title': 'Milliseconds (log)',
                              'type': 'log'},
                }

                fig_dict = {'data': data, 'layout': layout}

                pio.write_image(fig_dict,
                                root+'/'+matches.group(1)+".png",
                                engine="kaleido",
                                width=1600,
                                height=900,
                                validate=False)
                graphs[matches.group(1)] = root+'/'+matches.group(1)+".png"
                logger.info(f'Graph saved to: {graphs[matches.group(1)]}')
                os.remove('{}/{}.db'.format(root, matches.group(1)))
    return graphs


def show_graphs(directory, filename):
    regex = re.compile(r'(.*/)?[\w-]+_([\w-]+).(\w.+)')
    matches = regex.match(filename)
    if regex.match(filename) and matches.group(3) == 'json':
        # Initializes database for current file in current directory
        # Read by 20000 chunks
        disk_engine = create_engine(
            'sqlite:///{}.db'.format(matches.group(2)))

        j = 0
        index_start = 1
        chunk = 20000
        for df in pd.read_json(os.path.join(directory, filename),
                               lines=True,
                               chunksize=chunk):
            df.index += index_start

            columns = ['timestamp', 'latency', 'code']

            for c in df.columns:
                if c not in columns:
                    df = df.drop(c, axis=1)

            j += 1
            logger.info(f'completed {j*chunk} rows')

            df.to_sql('data', disk_engine, if_exists='append')
            index_start = df.index[-1] + 1

        df = pd.read_sql_query('SELECT * FROM data', disk_engine)

        data = [{
            'type': 'scatter',
            'x': df['timestamp'],
            'y': df['latency']/1000000,
            'mode': 'markers',
            'transforms': [
                {'type': 'groupby',
                 'groups': df['code'],
                 'color': 'code',
                 'styles': [
                    {'target': '0',
                     'value': {'marker': {'color': 'coral',
                                          'symbol': 'triangle-down'}}},
                    {'target': '200',
                     'value': {'marker': {'color': 'LightSkyBlue'}}},
                    {'target': '201',
                     'value': {'marker': {'color': 'LightSkyBlue'}}},
                    {'target': '400',
                     'value': {'marker': {'color': 'crimsom',
                                          'symbol': 'diamond'}}},
                    {'target': '500',
                     'value': {'marker': {'color': 'darkred',
                                          'symbol': 'diamond-tall'}}}]
                 }]
            }]

        layout = {
            'title': '<b>Latency per Request: {}</b>'.format(matches.group(2)),
            'xaxis': {'title': 'Time',
                      'showgrid': 'true',
                      'ticklabelmode': "period"},
            'yaxis': {'title': 'Milliseconds (log)', 'type': 'log'},
        }

        fig_dict = {'data': data, 'layout': layout}

        os.remove('{}.db'.format(matches.group(2)))

        pio.show(fig_dict,
                 engine="kaleido",
                 width=1600,
                 height=900,
                 validate=False)


def cma_graph(directory, filename):
    regex = re.compile(r'(.*/)?[\w-]+_([\w-]+).(\w.+)')
    matches = regex.match(filename)
    if regex.match(filename) and matches.group(3) == 'json':
        # Initializes database for current file in current directory
        # Read by 20000 chunks
        disk_engine = create_engine(
            'sqlite:///{}.db'.format(matches.group(2)))

        j = 0
        index_start = 1
        chunk = 20000
        for df in pd.read_json(os.path.join(directory, filename),
                               lines=True,
                               chunksize=chunk):
            df.index += index_start

            columns = ['timestamp', 'latency']

            for c in df.columns:
                if c not in columns:
                    df = df.drop(c, axis=1)

            j += 1
            logger.info(f'completed {j*chunk} rows')

            df.to_sql('data', disk_engine, if_exists='append')
            index_start = df.index[-1] + 1

        df = pd.read_sql_query('SELECT * FROM data', disk_engine)
        df_t = pd.DataFrame(df.iloc[:, -1])
        df_t.index = df.timestamp

        df_t['cma'] = df_t.expanding().mean()

        data = [{
            'type': 'line',
            'x': df_t.index,
            'y': df_t['cma']/1000000,
        }]

        layout = {
            'title': '<b>Cumulative AVG Latency: {}</b>'.format(
                matches.group(2)),
            'xaxis': {'title': 'Time',
                      'showgrid': 'true',
                      'ticklabelmode': "period"},
            'yaxis': {'title': 'Milliseconds (log)', 'type': 'linear'},
        }

        fig_dict = {'data': data, 'layout': layout}

        os.remove('{}.db'.format(matches.group(2)))

        pio.show(fig_dict,
                 engine="kaleido",
                 width=1600,
                 height=900,
                 validate=False)


def count_graph(directory, filename):
    regex = re.compile(r'(.*/)?[\w-]+_([\w-]+).(\w.+)')
    matches = regex.match(filename)
    if regex.match(filename) and matches.group(3) == 'json':
        # Initializes database for current file in current directory
        # Read by 20000 chunks
        disk_engine = create_engine(
            'sqlite:///{}.db'.format(matches.group(2)))

        j = 0
        index_start = 1
        chunk = 20000
        for df in pd.read_json(os.path.join(directory, filename),
                               lines=True,
                               chunksize=chunk):
            df.index += index_start

            columns = ['timestamp', 'latency']

            for c in df.columns:
                if c not in columns:
                    df = df.drop(c, axis=1)

            j += 1
            logger.info(f'completed {j*chunk} rows')

            df.to_sql('data', disk_engine, if_exists='append')
            index_start = df.index[-1] + 1

        df = pd.read_sql_query('SELECT * FROM data', disk_engine)
        df_t = pd.DataFrame(df.iloc[:, -1])
        df_t.index = df.timestamp

        df_t['count'] = df_t.expanding().count()

        data = [{
            'type': 'line',
            'x': df_t.index,
            'y': df_t['count'],
        }]

        layout = {
            'title': '<b>Request count : {}</b>'.format(matches.group(2)),
            'xaxis': {'title': 'Time',
                      'showgrid': 'true',
                      'ticklabelmode': "period"},
            'yaxis': {'title': 'Number of requests', 'type': 'linear'},
        }

        fig_dict = {'data': data, 'layout': layout}

        os.remove('{}.db'.format(matches.group(2)))

        pio.show(fig_dict,
                 engine="kaleido",
                 width=1600,
                 height=900,
                 validate=False)


def generate_summaries(directory):
    try:
        os.stat('{}/summaries'.format(directory))
    except FileNotFoundError:
        os.mkdir('{}/summaries'.format(directory))
    else:
        logger.error('Error with summaries folder.')
        exit(1)

    for root, _, files in os.walk(directory):
        for filename in files:
            regex = re.compile(r'([\w-]+)_([\w-]+).(\w.+)')
            matches = regex.match(filename)

            if 'summaries' not in root and regex.match(filename) and \
                    matches.group(3) == 'json':
                _summary_name = "{}/summaries/{}_{}-summary.json".format(
                               directory,
                               matches.group(1),
                               matches.group(2))
                logger.info(f'Generating summary for: {matches.group(2)}')
                subprocess.run(["vegeta", "report", "--type", "json",
                                "--output",
                                _summary_name,
                                "{}/{}".format(root, filename)])
                logger.info(f'Summary saved to: {_summary_name}')


def read_summaries(directory):
    summaries = {}
    # Walking through the folder to find all the summaries files
    # ignoring dirname since we don't need it in this implementation
    for root, _, files in os.walk(directory):
        for filename in files:
            regex = re.compile(
                r'[\w-]+_([\w-]+)-summary.(\w.+)')
            matches = regex.match(filename)

            if 'summaries' in root and regex.match(filename) and \
                    matches.group(2) == 'json':
                logger.info(f'Reading summary: {filename}')
                df = pd.read_json(root+'/'+filename, lines=True)

                lat = df['latencies'][0]
                summaries[matches.group(1)] = {
                    'requests':     int(df['requests']),
                    'rate':         float(df['rate']),
                    'duration':     int(df['duration']),
                    'min':          int(lat['min']),
                    'mean':         int(lat['mean']),
                    'max':          int(lat['max']),
                    'success':      float(df['success']),
                    'status_codes': df['status_codes'][0],
                    'errors':       df['errors'][0],
                }
    return summaries


def write_docx(directory, summaries, graphs, filename):
    date = datetime.datetime.utcnow()
    document = Document()

    document.add_heading('OCM Performance Test', 0)

    document.add_heading('Test # ', level=1)
    document.add_paragraph('Date: {}'.format(date.strftime("%Y-%m-%d")))

    document.add_heading('Description', level=2)
    document.add_paragraph('The purpose of this test is ...')

    document.add_heading('Notes', level=3)

    document.add_heading('Endpoints', level=2)

    table = document.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Enpoint'
    hdr_cells[1].text = 'Rate'
    hdr_cells[2].text = 'Notes'
    for r in summaries:
        row_cells = table.add_row().cells
        row_cells[0].text = r
        row_cells[1].text = '{:.2f}/s for {:.2f} minutes'.format(
            summaries[r]['rate'], summaries[r]['duration']/6e10)
        row_cells[2].text = ''

    document.add_heading('Per endpoint data', level=2)
    for r in summaries:
        document.add_heading('{}'.format(r), level=3)
        document.add_picture(graphs[r], width=Inches(16.6), height=Inches(9.4))
        p = document.add_paragraph(
            'Requests\t\tTotal: {}\t\tRate: {:.2f}\n'.format(
                summaries[r]['requests'], summaries[r]['rate']))
        p.add_run(
            'Duration\t\t{:.2f} minutes\n'.format(
                summaries[r]['duration']/6e10))
        p.add_run('Latencies\n')

        document.add_paragraph('Min: {:.4f} ms'.format(
            summaries[r]['min']/1e6), style='List Bullet')
        document.add_paragraph('Mean: {:.4f} ms'.format(
            summaries[r]['mean']/1e6), style='List Bullet')
        document.add_paragraph('Max: {:.4f} ms'.format(
            summaries[r]['max']/1e6), style='List Bullet')

        p2 = document.add_paragraph('Success\t\t{:.2f}%\n'.format(
            summaries[r]['success']*100))
        p2.add_run('Status Codes\t\t\n{}\n'.format(
            summaries[r]['status_codes']))
        p2.add_run('Error Set\t\t\n{}\n'.format(summaries[r]['errors']))
        p2.add_run('Notes').bold = True
        p2.add_run('\n')
        document.add_page_break()
    document.add_heading('Conclusion', level=2)
    document.add_paragraph('Make sure....', style='List Bullet')
    document.add_page_break()
    document.add_heading('Overall Screenshots', level=2)
    if '.docx' not in filename:
        filename = filename+'.docx'
    document.save('{}/{}'.format(directory, filename))


def upload_files(args):
    """
    Creates a folder and uploads the requests.tar.gz to it
    """
    uuid = ''
    tar_name = 'requests.tar.gz'
    tar = tarfile.open(os.path.join(args.directory, tar_name), "w|gz")
    for root, _, files in os.walk(args.directory):
        for filename in files:
            regex = re.compile(r'([\w-]+)_([\w-]+).(\w.+)')
            matches = regex.match(filename)
            if 'summaries' not in root and regex.match(filename) and \
                    matches.group(3) == 'json':
                uuid = matches.group(1)
                tar.add(os.path.join(args.directory, filename),
                        arcname='requests/{}'.format(filename))
                logger.info('Added file {} to archive'.format(
                            os.path.join(args.directory, filename)))
    tar.close()

    # Obtain current running diretory
    exec_folder = os.getcwd()

    # Change dir to folder where the tarball is
    # REST requests to SNAPPY need to be executed
    # from the folder where the file you need to upload is located
    os.chdir(args.directory)

    response = requests.post(f'{args.snappy_server}/auth/jwt/login',
                             data={
                                'password': args.snappy_password,
                                'username': args.snappy_user})
    if response.status_code != 200:
        logger.error(f'Authentication failed: {response.json()}')
        sys.exit(1)

    access_token = response.json()['access_token']
    auth_headers = {'Authorization': f'Bearer {access_token}'}

    query = {'filename': tar_name,
             'filedir': f'ocm/{uuid}'}
    response = requests.post(f'{args.snappy_server}/api',
                             params=query,
                             headers=auth_headers,
                             data=open(os.path.join(args.directory,
                                                    tar_name), 'rb'))
    if response.status_code != 200:
        logger.error(f'Upload of the file failed: {response.json()}')
        sys.exit(1)

    # Return to execution folder
    os.chdir(exec_folder)

    logger.info('File uploaded successfully')


def summarized_requests(path, index_name, test_id, test_name):
    """
    Yields a summarized request document for each line in a given Vegeta
    results file.

    The expected filename format is:
    40f696b8-0258-4a29-99f6-2767bd453548_create-cluster.json
    ^^^                                  ^^^
    Test UUID                            Test Name
    """

    for line in open(path, 'r'):
        req = json.loads(line)
        doc = {
            '_index': index_name,
            'test_name': test_name,
            'uuid': test_id,
            'timestamp': req['timestamp'],
            'code': req['code'],
            'method': req['method'],
            'url': req['url'],
            'latency_ns': req['latency'],
            'bytes_out': req['bytes_out'],
            'bytes_in': req['bytes_in'],
            'has_error': bool(req.get('error')),
            'has_body': bool(req.get('body')),
        }
        yield doc


def push_to_es(args):
    """
    The expected filename format is:
    40f696b8-0258-4a29-99f6-2767bd453548_create-cluster.json
    ^^^                                  ^^^
    Test UUID                            Test Name
    """
    # ElasticSearch Client
    es_host = os.getenv('ES')
    es_index = args.index
    assert es_host, "Did you forget to specify the environment variable `ES`?"
    es = Elasticsearch(es_host, use_ssl=False, verify_certs=False)
    logger.info('Connected to ElasticSearch')
    es.indices.create(index=es_index, ignore=400)  # Ignore IndexAlreadyExists

    for root, _, files in os.walk(args.directory):
        for filename in files:
            regex = re.compile(r'([\w-]+)_([\w-]+).(\w.+)')
            matches = regex.match(filename)
            if 'summaries' not in root and regex.match(filename) and \
                    matches.group(3) == 'json':
                test_id = matches.group(1)
                test_name = matches.group(2)

                logger.info("Indexing result file: %s" % filename)
                helpers.bulk(es, summarized_requests(os.path.join(root,
                                                                  filename),
                                                     es_index,
                                                     test_id,
                                                     test_name))


def main():
    """Automation script to process the results of a test
        - Generate latency spread graphs
        - Generate vegita reports
        - Generate full report
        - Upload results files
    """

    date = datetime.datetime.utcnow()
    parent_parser = argparse.ArgumentParser(add_help=False)
    parent_parser.add_argument('--dir',
                               dest="directory",
                               default='.',
                               required=True,
                               help='directory path were results are stored')
    parent_parser.add_argument('--debug',
                               default=False,
                               required=False,
                               action='store_true',
                               dest="debug",
                               help='debug flag')

    main_parser = argparse.ArgumentParser()

    action_subparsers = main_parser.add_subparsers(title="action",
                                                   dest="action_command")

    graph_parser = action_subparsers.add_parser("graph",
                                                help="generate the graps \
                                                 for the results file",
                                                parents=[parent_parser])

    graph_parser.add_argument('--filename',
                              dest="filename",
                              help='filename of a result to display the graph. \
                                (Overrides generating all graphs.)')

    cma_parser = action_subparsers.add_parser("cma",
                                              help="generate cummulative average graph \
                                               for the results file",
                                              parents=[parent_parser])

    cma_parser.add_argument('--filename',
                            dest="filename",
                            help='filename of a result to display the graph',
                            required=True)

    count_parser = action_subparsers.add_parser("count",
                                                help="generate cummulative count of requests graph \
                                                for the results file",
                                                parents=[parent_parser])

    count_parser.add_argument('--filename',
                              dest="filename",
                              help='filename of a result to display the graph',
                              required=True)

    action_subparsers.add_parser("summary",
                                 help="generates vegeta \
                                  summary for results",
                                 parents=[parent_parser])

    report_parser = action_subparsers.add_parser("report",
                                                 help="generates report",
                                                 parents=[parent_parser])

    report_parser.add_argument('--filename',
                               dest='filename',
                               default='report-{}.docx'.format(
                                    date.strftime("%Y-%m-%d")),
                               help='name for the report file.')

    upload_parser = action_subparsers.add_parser("upload",
                                                 help="uploads test results",
                                                 parents=[parent_parser])

    upload_parser.add_argument('--server',
                               dest="snappy_server",
                               help='Snappy server URL',
                               required=True)

    upload_parser.add_argument('--user',
                               dest="snappy_user",
                               help='User for authenticating to snappy',
                               required=True)

    upload_parser.add_argument('--password',
                               dest="snappy_password",
                               help='Password for authenticating to snappy',
                               required=True)

    es_bulk = action_subparsers.add_parser("esbulk",
                                           help="uploads results to ES",
                                           parents=[parent_parser])

    es_bulk.add_argument('--index',
                         dest="index",
                         help='ES index where the documents will be stored.')

    args = main_parser.parse_args()

    if args.action_command == 'graph':
        if args.filename is not None:
            show_graphs(args.directory, args.filename)
        else:
            generate_graphs(args.directory)
    elif args.action_command == 'cma':
        cma_graph(args.directory, args.filename)
    elif args.action_command == 'count':
        count_graph(args.directory, args.filename)
    elif args.action_command == 'summary':
        generate_summaries(args.directory)
    elif args.action_command == 'report':
        graphs = generate_graphs(args.directory)
        generate_summaries(args.directory)
        summaries = read_summaries(args.directory)
        write_docx(args.directory, summaries, graphs, args.filename)
    elif args.action_command == 'upload':
        upload_files(args)
    elif args.action_command == 'esbulk':
        push_to_es(args)


if __name__ == "__main__":
    main()
