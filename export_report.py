#!/usr/bin/python3

import plotly.io as pio
import pandas as pd
from sqlalchemy import create_engine
from docx import Document
from docx.shared import Inches
import os
import re
import subprocess
import sys
import datetime


def generate_graphs(directory):
    graphs = {}
    # Walking through the folder to find all the result files
    # ignoring dirname since we don't need it in this implementation
    for root, _, files in os.walk(directory):
        for filename in files:
            regex = re.compile(r'[\w-]+_([\w-]+).(\w.+)')
            matches = regex.match(filename)

            if 'reports' not in root and regex.match(filename) and \
                    matches.group(2) == 'json':
                print('Generating graph for: {}'.format(matches.group(1)))

                # Initializes database for current file in current directory
                # Read by 20000 chunks
                disk_engine = create_engine(
                    'sqlite:///{}/{}.db'.format(root, matches.group(1)))

                j = 0
                index_start = 1
                chunk = 20000
                for df in pd.read_json(root+'/'+filename,
                                       lines=True,
                                       chunksize=chunk):
                    df.index += index_start

                    columns = ['timestamp', 'latency', 'code']

                    for c in df.columns:
                        if c not in columns:
                            df = df.drop(c, axis=1)

                    j += 1
                    print('completed {} rows'.format(j*chunk))

                    df.to_sql('data', disk_engine, if_exists='append')
                    index_start = df.index[-1] + 1

                df = pd.read_sql_query('SELECT * FROM data', disk_engine)

                data = [{
                    'type': 'scatter',
                    'x': df['timestamp'],
                    'y': df['latency']/1000000,
                    'mode': 'markers',
                    'transforms': [{
                        'type': 'groupby',
                        'groups': df['code']
                    }]
                }]

                layout = {
                    'title': '<b>Latency per Request: {}</b>'.format(
                        matches.group(1)),
                    'xaxis': {'title': 'Time',
                              'showgrid': 'true',
                              'ticklabelmode': "period"},
                    'yaxis': {'title': 'Milliseconds (log)',
                              'type': 'log'},
                }

                fig_dict = {'data': data, 'layout': layout}

                pio.write_image(fig_dict,
                                root+'/'+matches.group(1)+".png",
                                engine="kaleido",
                                width=1600,
                                height=900,
                                validate=False)
                graphs[matches.group(1)] = root+'/'+matches.group(1)+".png"
                print('Graph saved to: {}'.format(graphs[matches.group(1)]))
                os.remove('{}/{}.db'.format(root, matches.group(1)))
    return graphs


def show_graphs(filename):
    regex = re.compile(r'(.*/)?[\w-]+_([\w-]+).(\w.+)')
    matches = regex.match(filename)
    if regex.match(filename) and matches.group(3) == 'json':
        # Initializes database for current file in current directory
        # Read by 20000 chunks
        disk_engine = create_engine(
            'sqlite:///{}.db'.format(matches.group(2)))

        j = 0
        index_start = 1
        chunk = 20000
        for df in pd.read_json(filename,
                               lines=True,
                               chunksize=chunk):
            df.index += index_start

            columns = ['timestamp', 'latency', 'code']

            for c in df.columns:
                if c not in columns:
                    df = df.drop(c, axis=1)

            j += 1
            print('completed {} rows'.format(j*chunk))

            df.to_sql('data', disk_engine, if_exists='append')
            index_start = df.index[-1] + 1

        df = pd.read_sql_query('SELECT * FROM data', disk_engine)

        data = [{
            'type': 'scatter',
            'x': df['timestamp'],
            'y': df['latency']/1000000,
            'mode': 'markers',
            'transforms': [
                {'type': 'groupby',
                 'groups': df['code']}]
                }]

        layout = {
            'title': '<b>Latency per Request: {}</b>'.format(matches.group(2)),
            'xaxis': {'title': 'Time',
                      'showgrid': 'true',
                      'ticklabelmode': "period"},
            'yaxis': {'title': 'Milliseconds (log)', 'type': 'log'},
        }

        fig_dict = {'data': data, 'layout': layout}

        os.remove('{}.db'.format(matches.group(2)))

        pio.show(fig_dict,
                 engine="kaleido",
                 width=1600,
                 height=900,
                 validate=False)


def generate_reports(directory):
    if not os.stat('{}/reports'.format(directory)):
        os.mkdir('{}/reports'.format(directory))
    for root, _, files in os.walk(directory):
        for filename in files:
            regex = re.compile(r'([\w-]+)_([\w-]+).(\w.+)')
            matches = regex.match(filename)

            if 'reports' not in root and regex.match(filename) and \
                    matches.group(3) == 'json':
                _report_name = "{}/reports/{}_{}-report.json".format(
                               directory,
                               matches.group(1),
                               matches.group(2))
                print('Generating report for: {}'.format(matches.group(2)))
                subprocess.run(["vegeta", "report", "--type", "json",
                                "--output",
                                _report_name,
                                "{}/{}".format(root, filename)])
                print('Report saved to: {}'.format(_report_name))


def read_reports(directory):
    reports = {}
    # Walking through the folder to find all the report files
    # ignoring dirname since we don't need it in this implementation
    for root, _, files in os.walk(directory):
        for filename in files:
            regex = re.compile(
                r'[\w-]+_([\w-]+)-report.(\w.+)')
            matches = regex.match(filename)

            if 'reports' in root and regex.match(filename) and \
                    matches.group(2) == 'json':
                print('Reading report: {}'.format(filename))
                df = pd.read_json(root+'/'+filename, lines=True)

                lat = df['latencies'][0]
                reports[matches.group(1)] = {
                    'requests':     int(df['requests']),
                    'rate':         float(df['rate']),
                    'duration':     int(df['duration']),
                    'min':          int(lat['min']),
                    'mean':         int(lat['mean']),
                    'max':          int(lat['max']),
                    'success':      float(df['success']),
                    'status_codes': df['status_codes'][0],
                    'errors':       df['errors'][0],
                }
    return reports


def write_docx(directory, reports, graphs):
    date = datetime.datetime.utcnow()
    document = Document()

    document.add_heading('OCM Performance Test', 0)

    document.add_heading('Test # ', level=1)
    document.add_paragraph('Date: {}'.format(date.strftime("%Y-%m-%d")))

    document.add_heading('Description', level=2)
    document.add_paragraph('The purpose of this test is ...')

    document.add_heading('Notes', level=3)

    document.add_heading('Endpoints', level=2)

    table = document.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Enpoint'
    hdr_cells[1].text = 'Rate'
    hdr_cells[2].text = 'Notes'
    for r in reports:
        row_cells = table.add_row().cells
        row_cells[0].text = r
        row_cells[1].text = '{:.2f}/s for {:.2f} minutes'.format(
            reports[r]['rate'], reports[r]['duration']/6e10)
        row_cells[2].text = ''

    document.add_heading('Per endpoint data', level=2)
    for r in reports:
        document.add_heading('{}'.format(r), level=3)
        document.add_picture(graphs[r], width=Inches(16.6), height=Inches(9.4))
        p = document.add_paragraph(
            'Requests\t\tTotal: {}\t\tRate: {:.2f}\n'.format(
                reports[r]['requests'], reports[r]['rate']))
        p.add_run(
            'Duration\t\t{:.2f} minutes\n'.format(
                reports[r]['duration']/6e10))
        p.add_run('Latencies\n')

        document.add_paragraph('Min: {:.4f} ms'.format(
            reports[r]['min']/1e6), style='List Bullet')
        document.add_paragraph('Mean: {:.4f} ms'.format(
            reports[r]['mean']/1e6), style='List Bullet')
        document.add_paragraph('Max: {:.4f} ms'.format(
            reports[r]['max']/1e6), style='List Bullet')

        p2 = document.add_paragraph('Success\t\t{:.2f}%\n'.format(
            reports[r]['success']*100))
        p2.add_run('Status Codes\t\t\n{}\n'.format(reports[r]['status_codes']))
        p2.add_run('Error Set\t\t\n{}\n'.format(reports[r]['errors']))
        p2.add_run('Notes').bold = True
        p2.add_run('\n')
        document.add_page_break()
    document.add_heading('Conclusion', level=2)
    document.add_paragraph('Make sure....', style='List Bullet')
    document.add_page_break()
    document.add_heading('Overall Screenshots', level=2)
    document.save('{}/report-{}.docx'.format(directory,
                                             date.strftime("%Y-%m-%d")))


def main():
    directory = "."
    args = sys.argv[1:]
    if len(args) > 0:
        if args[0] == 'graph':
            show_graphs(args[1])
            exit()
        elif args[0] == 'report':
            generate_reports(args[1])
            exit()
        directory = args[0]

    graphs = generate_graphs(directory)
    generate_reports(directory)
    reports = read_reports(directory)
    write_docx(directory, reports, graphs)


if __name__ == "__main__":
    main()
